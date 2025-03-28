from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


class DocxController:
    def __init__(self):
        self.document = Document()

    def add_heading(self, text, level=1):
        self.document.add_heading(text, level=level)

    def add_paragraph(self, text, font_name="맑은 고딕", font_size=12, bold=False):
        paragraph = self.document.add_paragraph()
        run = paragraph.add_run(text)
        font = run.font
        font.name = font_name
        font.size = Pt(font_size)
        font.bold = bold

        # 폰트 설정 안정화 (한글 폰트 호환성 강화)
        r = run._element
        r.rPr.rFonts.set('eastAsia', font_name)

        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    def save(self, file_path):
        self.document.save(file_path)
        return True
